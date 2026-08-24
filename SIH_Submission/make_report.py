"""Generate UdaanSetu SIH 2026 project report."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x0C, 0x3B, 0x26)
GREEN = RGBColor(0x16, 0xA3, 0x4A)

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def bullet(text, bold_head=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_head:
        r = p.add_run(bold_head)
        r.bold = True
        p.add_run(' — ' + text)
    else:
        p.add_run(text)
    return p


# ---------------- Cover ----------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('UdaanSetu')
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = DARK
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Bridge to Flight — Research to Impact')
r.font.size = Pt(16)
r.font.color.rgb = GREEN
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Innovation Ecosystem Platform\n')
r.font.size = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Smart India Hackathon 2026  •  Problem ID: SIH1608\n')
r.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[TEAM NAME]  •  [INSTITUTE NAME]  •  [CITY, STATE]\n')
r.font.size = Pt(12)
doc.add_page_break()

# ---------------- 1. Executive Summary ----------------
h1('1. Executive Summary')
para('UdaanSetu is a unified innovation ecosystem platform that guides a project from '
     'research through IPR, mentorship, funding, incubation, and startup launch — all on one '
     'platform. It solves the fundamental fragmentation problem in India\u2019s innovation landscape: '
     'world-class research rarely becomes market-ready products because innovators lack '
     'guidance on the next step, visibility into government schemes, transparent IPR tracking, '
     'and access to mentors and funding.')
para('The platform is production-grade: full-stack Dockerized application, 60+ REST APIs, '
     'role-based access for all five ecosystem stakeholders, ML-based risk prediction and '
     'semantic matching, government API integrations (demo mode), and 190+ automated tests.')

# ---------------- 2. Problem Statement ----------------
h1('2. Problem Statement')
para('India publishes 50,000+ research papers every year, yet fewer than 500 ever become '
     'startups. The journey from laboratory to market is fragmented:', )
bullet('Research results stay in journals and never move toward products')
bullet('Innovators lack guidance on what to do next (patent? funding? pilot?)')
bullet('Government schemes and grants are invisible to the innovators who need them')
bullet('IPR filing is opaque, slow, and difficult to track')
bullet('Mentors, incubators, and investors cannot see the ecosystem pipeline')
para('The result: billions in lost innovation potential and an average lab-to-market time of '
     '3\u20135 years.', bold=True)

# ---------------- 3. Proposed Solution ----------------
h1('3. Proposed Solution')
para('UdaanSetu unifies the entire innovation lifecycle on a single platform. Every project is '
     'guided through six stages with active support:', )
for sname, desc in [
    ('Research', 'Researchers register projects with stage, district, sector, funding need and milestones.'),
    ('Innovation', 'Semantic duplicate detection prevents redundancy; AI recommends the next step.'),
    ('IPR / Patent', 'Track patent filings, application numbers, stages, and due dates transparently.'),
    ('Mentor / Funding / Incubator', 'Smart matching connects each innovator to the right mentors, schemes, and incubators based on semantic similarity.'),
    ('Startup', 'Launch with jobs-created, farmers-reached, and revenue tracked.'),
    ('Impact', 'Measure sector and district-level outcomes to show ecosystem value.'),
]:
    bullet(desc, bold_head=sname)

# ---------------- 4. Key Features ----------------
h1('4. Key Features')
for name, desc in [
    ('Role-Based Dashboards', 'Admin, Researcher, Mentor, Investor, and Incubator each get a tailored view of the ecosystem.'),
    ('ML Risk Prediction', 'A trained gradient-boosting model scores every project (0\u2013100) with explainable reasons such as overdue milestones, low progress, or risky stage.'),
    ('Semantic Search & Smart Matching', 'sentence-transformer embeddings match innovators with mentors, schemes, and incubators by meaning, not just keywords.'),
    ('Duplicate Detection', 'NLP clustering flags near-duplicate research and IP to keep the ecosystem clean.'),
    ('Government Integrations', 'Aadhaar eKYC, DigiLocker, Startup India, IP India, and ONDC flows (mock data, demo mode).'),
    ('Security & Compliance', 'JWT authentication, Argon2 password hashing, role-based access control, input sanitization, rate limiting, and a full audit log.'),
]:
    bullet(desc, bold_head=name)

# ---------------- 5. Technology Stack ----------------
h1('5. Technology Stack')
for name, desc in [
    ('Frontend', 'Next.js 15, React 19, TypeScript, custom accessible design system (WCAG 2.1 AA).'),
    ('Backend', 'FastAPI, Python 3.12, SQLAlchemy 2.0, Pydantic v2, Uvicorn.'),
    ('Database', 'PostgreSQL with Redis caching (in-memory fallback).'),
    ('ML / AI', 'scikit-learn (gradient boosting, clustering, TF-IDF), sentence-transformers embeddings, retrain-on-real-data pipeline.'),
    ('Security', 'PyJWT, pwdlib/Argon2, RBAC, rate limiting, audit logging.'),
    ('Infrastructure', 'Docker Compose, health checks, automated test suites (172 backend + 19 frontend).'),
]:
    bullet(desc, bold_head=name)

# ---------------- 6. Architecture ----------------
h1('6. Architecture')
para('Stateless FastAPI backend exposes 60+ REST endpoints consumed by the Next.js frontend. '
     'The ML engine provides risk prediction, semantic search, matching, and duplicate detection. '
     'PostgreSQL stores normalized relational data; Redis caches dashboards and analytics. '
     'All services run under Docker Compose with health checks. Admin-only endpoints trigger '
     'model retraining on real database records, with a model registry for versioning and '
     'promotion.')

# ---------------- 7. AI / ML Approach ----------------
h1('7. AI / ML Approach')
bullet('Risk Prediction: gradient-boosting classifier trained on project features (progress, milestones, stage, funding ratio, days active). Pseudo-labels are derived from domain heuristics; falls back to synthetic training when real samples are sparse.')
bullet('Semantic Matching: sentence-transformer embeddings (all-MiniLM-L6-v2) with TF-IDF fallback for offline environments.')
bullet('Duplicate Detection: agglomerative clustering on cosine distances of embeddings.')
bullet('Retraining: when ≥20 real records exist, the pipeline trains on real data and registers a new model version; predictions remain explainable.')

# ---------------- 8. Security ----------------
h1('8. Security')
bullet('JWT access tokens with refresh flow; Argon2 password hashing')
bullet('Role-based access control across all five roles')
bullet('Input sanitization on all text fields')
bullet('Rate limiting (120 requests/min/IP) with health-endpoint exemption')
bullet('Full audit log of create / update / delete actions')
bullet('Token blacklisting on logout')

# ---------------- 9. Demo Script Summary ----------------
h1('9. Demo Walkthrough')
for t, title, desc in [
    ('00:00', 'Login as Researcher', 'admin@udaansetu.demo / Demo@123'),
    ('00:30', 'Create Research Project', 'set stage, district, sector, funding need'),
    ('01:15', 'Add Milestones', 'track progress; overdue flagged red'),
    ('01:45', 'AI Risk Assessment', 'score + reasons; at-risk project highlighted'),
    ('02:15', 'Smart Recommendations', 'matching mentors, schemes, incubators'),
    ('02:45', 'IPR Filing Tracker', 'patent stages & application status'),
    ('03:30', 'Govt Integrations', 'Aadhaar / DigiLocker / Startup India / ONDC'),
    ('04:00', 'Analytics & Impact', 'dashboard metrics, sector/district impact'),
    ('04:45', 'Audit & Security', 'role-based access, audit log, rate limiting'),
]:
    bullet(f'{t} — {title}: {desc}')

# ---------------- 10. Roadmap ----------------
h1('10. Roadmap')
for phase, desc in [
    ('Phase 1 (Done)', 'Core platform, 60+ APIs, ML pipeline, government integrations, tests, Dockerization.'),
    ('Phase 2 (Next)', 'Ingest real research / patent / scheme datasets and retrain models on real data.'),
    ('Phase 3', 'Production deployment, CI/CD, monitoring, model drift detection.'),
    ('Phase 4', 'Live government API connections, mobile app, multi-language support.'),
]:
    bullet(desc, bold_head=phase)

# ---------------- 11. Team ----------------
h1('11. Team')
p = doc.add_paragraph()
r = p.add_run('[TEAM NAME] — Team Leader')
r.bold = True
for role in ['[Name] — [Role]', '[Name] — [Role]', '[Name] — [Role]',
             '[Name] — [Role]', '[Name] — [Role]']:
    doc.add_paragraph(role, style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Institute: [INSTITUTE NAME]')
r.bold = True
doc.add_paragraph('City/State: [CITY, STATE]')

doc.save(r"C:\Users\Rudra\Desktop\UdaanSetu\SIH_Submission\UdaanSetu_SIH2026_Report.docx")
print("Report saved: SIH_Submission/UdaanSetu_SIH2026_Report.docx")